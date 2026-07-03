from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseForbidden, FileResponse
from django.contrib import messages
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Tugas, Logbook
from .forms import LogbookForm
from accounts.utils import catat_log

@login_required
def daftar_tugas(request):
    if request.method == 'POST':
        tugas_id = request.POST.get('tugas_id')
        if tugas_id:
            try:
                tugas = Tugas.objects.get(id=tugas_id, mahasiswa=request.user)
                tugas.status = 'SELESAI'
                tugas.save()
                messages.success(request, 'Tugas berhasil ditandai selesai.')
            except Tugas.DoesNotExist:
                messages.error(request, 'Tugas tidak ditemukan.')
        return redirect('daftar_tugas')

    from django.utils import timezone
    tugas_list = Tugas.objects.filter(mahasiswa=request.user).order_by('-deadline')
    context = {
        'page_title': 'Tugas Saya',
        'tugas_list': tugas_list,
        'total': len(tugas_list),
        'today': timezone.localdate(),
    }
    return render(request, 'user/user_tugas.html', context)

@login_required
def entri_logbook(request):
    logbook_id = request.GET.get('id')
    logbook = None
    
    if logbook_id:
        logbook = get_object_or_404(Logbook, pk=logbook_id)
        
        # Keamanan: Pastikan hanya pemilik logbook yang bisa mengedit
        if logbook.user != request.user:
            return HttpResponseForbidden("Anda tidak memiliki akses untuk mengedit logbook ini.")
            
        # Keamanan: Pastikan statusnya REVISI
        if logbook.status != 'REVISI':
            messages.error(request, "Hanya logbook dengan status REVISI yang bisa diubah.")
            return redirect('entri_logbook')

    # Logika Form (POST/GET)
    if request.method == 'POST':
        form = LogbookForm(request.POST, request.FILES, instance=logbook)
        if form.is_valid():
            logbook_instance = form.save(commit=False)
            # Jika ini data baru, set user-nya
            if not getattr(logbook_instance, 'user_id', None):
                logbook_instance.user = request.user
            
            # Setiap perubahan atau entri baru, statusnya menunggu
            logbook_instance.status = 'MENUNGGU' 
            logbook_instance.save()
            catat_log(request.user, "Mengirimkan laporan logbook harian", "CREATE")
            messages.success(request, 'Logbook berhasil disimpan dan sedang menunggu persetujuan.')
            return redirect('dashboard')
    else:
        form = LogbookForm(instance=logbook)

    # Data riwayat aktivitas asli
    logbook_list = Logbook.objects.filter(user=request.user).order_by('-tanggal')
    
    context = {
        'page_title': 'Perbaikan Logbook' if logbook else 'Entri Logbook',
        'form': form,
        'logbook': logbook,
        'is_edit_mode': logbook is not None,
        'logbook_list': logbook_list,
        'total': len(logbook_list),
    }
    return render(request, 'user/user_logbook.html', context)

@login_required
def riwayat_logbook_peserta(request):
    logbooks = Logbook.objects.filter(user=request.user).order_by('-tanggal')
    return render(request, 'user/riwayat_logbook.html', {'logbooks': logbooks})

@login_required
def ekspor_logbook_word(request):
    logbooks = Logbook.objects.filter(user=request.user).order_by('tanggal')
    
    doc = Document()
    
    # Judul Dokumen
    judul = doc.add_paragraph('Format Logbook Kegiatan Mahasiswa\nLogbook Harian Kegiatan Magang')
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informasi Mitra dan Mahasiswa
    info_table = doc.add_table(rows=4, cols=3)
    info_data = [
        ('Nama Mitra', ':', 'Dinas Komunikasi, Informatika, Persandian dan Statistik Kabupaten Bekasi'),
        ('Alamat Mitra', ':', 'Desa Sukamahi, Kecamatan Cikarang Pusat, Kabupaten Bekasi, Jawa Barat 17811'),
        ('Nama Mahasiswa', ':', getattr(request.user, 'first_name', '')),
        ('NIM Mahasiswa', ':', getattr(request.user, 'nim', ''))
    ]
    for i, row in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]

    doc.add_paragraph()
    
    # Tabel Utama Logbook
    tabel_logbook = doc.add_table(rows=1, cols=6)
    tabel_logbook.style = 'Table Grid'
    hdr_cells = tabel_logbook.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Tanggal/Waktu'
    hdr_cells[2].text = 'Kegiatan'
    hdr_cells[3].text = 'Uraian Kegiatan'
    hdr_cells[4].text = 'Paraf'
    hdr_cells[5].text = 'Keterangan'
    
    for index, log in enumerate(logbooks):
        row_cells = tabel_logbook.add_row().cells
        row_cells[0].text = str(index + 1)
        row_cells[1].text = str(log.tanggal)
        row_cells[2].text = str(log.judul)
        row_cells[3].text = str(log.detail_pekerjaan)
        row_cells[4].text = '' 
        row_cells[5].text = '' 

    doc.add_paragraph()
    
    # Bagian Tanda Tangan
    ttd_table = doc.add_table(rows=4, cols=2)
    ttd_cells = ttd_table.rows[0].cells
    ttd_cells[0].text = 'Menyetujui,\nDosen Pembimbing Lapangan'
    ttd_cells[1].text = 'Bekasi ,\nDosen Pembimbing'
    
    doc.add_paragraph('\n\n')
    
    ttd_cells_bawah = ttd_table.rows[3].cells
    ttd_cells_bawah[0].text = 'NIP'
    ttd_cells_bawah[1].text = 'NIDN'

    # Proses Ekspor
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return FileResponse(buffer, as_attachment=True, filename=f'Logbook_{getattr(request.user, "nim", "Peserta")}.docx')